import lab as l 
from tkinter import messagebox 
import pyodbc
from docx import Document
from docx.shared import Pt

def write_taqreer() :
    
    doc = Document()

    connection_string = f"""Driver={{SQL Server}};
        Server={"DESKTOP-CN9E67A"};
        Database={"lab"};
        Trusted_connection=yes;
        UID={"DESKTOP-CN9E67A"};
        PWD={"ahmed2003"}"""
    conn = pyodbc.connect(connection_string)
    cur = conn.cursor()
        
    cur.execute("select s_name from patients where lr_ref = ?", (l.en41.get()))
    rows = cur.fetchall()
         
    try :
            
        if rows ==  []:
                
            l.en41.delete(0,l.END)
            cur.close()
            raise ValueError
        
        else :
    
            p = doc.add_paragraph()
            p2 = doc.add_paragraph()
            p3 = doc.add_paragraph()
            p4 = doc.add_paragraph()
            p5 = doc.add_paragraph()
            p6 = doc.add_paragraph()
            p7 = doc.add_paragraph()
            p8 = doc.add_paragraph()
            p9 = doc.add_paragraph()
            p10 = doc.add_paragraph()
            p11 = doc.add_paragraph()
            p12 = doc.add_paragraph()
    
            p.add_run(f"        DEAR PROF. DR/{rows[0][0]}")

            for run in p.runs:
                run.font.size = Pt(18)
                run.font.bold = True
        
            cur.execute("select name from patients where lr_ref = ?", (l.en41.get()))
            rows = cur.fetchall()    
            p2.add_run(f"Patient name : {rows[0][0]}")
            p2.add_run(f"                                          For lab use")
    
            cur.execute("select age from patients where lr_ref = ?", (l.en41.get()))
            rows = cur.fetchall() 
    
            p2.add_run(f"\nAge : {rows[0][0]}")
    
            cur.execute("select sex from patients where lr_ref = ?", (l.en41.get()))
            rows = cur.fetchall() 
    
            p2.add_run(f"             sex : {rows[0][0]}")
    
            cur.execute("select lr_ref from patients where lr_ref = ?", (l.en41.get()))
            rows = cur.fetchall()  
    
            p2.add_run(f"                                Lab.Reference no. : {rows[0][0]}")
    
            cur.execute("select rc_date from patients where lr_ref = ?", (l.en41.get()))
            rows = cur.fetchall() 
    
            p2.add_run(f"\nReceiving date : {rows[0][0]}")
    
            cur.execute("select pre_ref from patients where lr_ref = ?", (l.en41.get()))
            rows = cur.fetchall() 
    
            p2.add_run(f"                                Previous biopsy no. : 147/23")
    
            cur.execute("select rp_date from patients where lr_ref = ?", (l.en41.get()))
            rows = cur.fetchall() 
    
            p2.add_run(f"\nReporting date : {rows[0][0]}")

            for run in p2.runs:
                run.font.size = Pt(14) 
        
        
            p3.add_run(f"Nature of specimen : ")
    
            for run in p3.runs:
                run.font.size = Pt(16)
                run.font.underline = True
                run.font.bold = True
    
            cur.execute("select code from patients where lr_ref = ?", (l.en41.get()))
            rows = cur.fetchall()
            cur.execute(f"select nature from diagnosis where code = '{rows[0][0]}'")
            rows = cur.fetchall()
            p4.add_run(f"{rows[0][0]}")
    
            for run in p4.runs:
                run.font.size = Pt(14) 
                run.font.name = "Bodoni MT"
        
    
            p5.add_run(f"Gross : ")
    
            for run in p5.runs:
                run.font.size = Pt(16)
                run.font.underline = True 
                run.font.bold = True   
        
            cur.execute("select code from patients where lr_ref = ?", (l.en41.get()))
            rows = cur.fetchall()
            cur.execute(f"select gross from diagnosis where code = '{rows[0][0]}'")
            rows = cur.fetchall()
            p6.add_run(f"{rows[0][0]}")
    
            for run in p6.runs:
                run.font.size = Pt(14) 
                run.font.name = "Arial"    
        
    
            p7.add_run(f"Microscopic : ")
    
            for run in p7.runs:
                run.font.size = Pt(16)
                run.font.underline = True 
                run.font.bold = True   
        
            cur.execute("select code from patients where lr_ref = ?", (l.en41.get()))
            rows = cur.fetchall()
            cur.execute(f"select microscope from diagnosis where code = '{rows[0][0]}'")
            rows = cur.fetchall()
            p8.add_run(f"{rows[0][0]}")
    
            for run in p8.runs:
                run.font.size = Pt(14) 
                run.font.name = "Arial" 
        
    
            p9.add_run(f"Conclusion : ")
    
            for run in p9.runs:
                run.font.size = Pt(18)
                run.font.underline = True
                run.font.bold = True   
        
            cur.execute("select code from patients where lr_ref = ?", (l.en41.get()))
            rows = cur.fetchall()
            cur.execute(f"select conclusion from diagnosis where code = '{rows[0][0]}'")
            rows = cur.fetchall()
            p10.add_run(f"{rows[0][0]}")
    
            for run in p10.runs:
                run.font.name = "Arial"
                run.font.size = Pt(16)
                run.font.bold = True  
        
    
            p11.add_run("                                                                      SIGNATURE\n")
            p11.add_run("                                                                   Dr.Rehab Essia\n")
            p11.add_run("                                                              Dr.Eman El-Shenawy\n")
            p11.add_run("                                                                   Dr.Nesma Salah")
    
            for run in p11.runs:
                run.font.size = Pt(16)
                run.font.bold = True
                run.font.name = "Bodoni MT"  
        
    
            cur.execute("select shareeha from patients where lr_ref = ?", (l.en41.get()))
            rows = cur.fetchall()
    
            p12.add_run(f"تم تسليم المريض عدد {rows[0][0]} شرائح")
    
            for run in p12.runs:
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.underline = True   

            cur.execute("select name from patients where lr_ref = ?", (l.en41.get()))
            rows = cur.fetchall()
            doc.save(f"D:\\OneDrive\\Desktop\\My work\\Lab-Project\\reports\\{rows[0][0]}.docx")
            
            l.en41.delete(0,l.END)
            messagebox.showinfo("Info", "The file saved sucsessfylly!")
            
    except :
            
          messagebox.showwarning("Warning", "This reference is not correct")        